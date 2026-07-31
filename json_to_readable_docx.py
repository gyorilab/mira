import json
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


INPUT_FOLDER = Path(
    "../miradb/table_extraction_eval/results/structured_outputs_v2"
)

OUTPUT_FILE = Path(
    "../miradb/table_extraction_eval/results/"
    "Structured_Parameter_Extraction_Report.docx"
)


def safe_text(value):
    if value is None:
        return "-"
    if isinstance(value, list):
        return ", ".join(str(item) for item in value)
    return str(value)


def evidence_text(evidence):
    if not evidence:
        return "-"

    entries = []
    for item in evidence:
        table = item.get("source_table", "-")
        row = item.get("source_row", "-")
        column = item.get("source_column")

        text = f"{table}, row {row}"
        if column:
            text += f", column {column}"

        entries.append(text)

    return "; ".join(entries)


def set_cell_text(cell, text, bold=False, font_size=8):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(safe_text(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_summary(document, data):
    document.add_paragraph(
        f"Contains parameter information: "
        f"{data.get('contains_parameter_information', False)}"
    )

    parameters = data.get("parameters", [])

    parameter_count = sum(
        p.get("value_type") == "parameter"
        for p in parameters
    )
    initial_count = sum(
        p.get("value_type") == "initial_condition"
        for p in parameters
    )
    definition_count = sum(
        p.get("value_type") == "definition_only"
        for p in parameters
    )

    document.add_paragraph(f"Total extracted records: {len(parameters)}")
    document.add_paragraph(f"Model parameters: {parameter_count}")
    document.add_paragraph(f"Initial conditions: {initial_count}")
    document.add_paragraph(f"Definition-only records: {definition_count}")

    notes = data.get("notes")
    if notes:
        document.add_paragraph(f"Notes: {notes}")


def add_parameter_table(document, parameters):
    headers = [
        "No.",
        "Symbol",
        "Parameter name",
        "Description",
        "Value",
        "Unit",
        "Range",
        "Uncertainty",
        "Prior",
        "Source",
        "Type",
        "Context",
        "Evidence",
    ]

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
            font_size=8,
        )

    for number, parameter in enumerate(parameters, start=1):
        row = table.add_row().cells

        uncertainty = safe_text(parameter.get("uncertainty"))
        uncertainty_type = parameter.get("uncertainty_type")

        if uncertainty != "-" and uncertainty_type:
            uncertainty = f"{uncertainty_type}: {uncertainty}"

        values = [
            number,
            parameter.get("symbol"),
            parameter.get("parameter_name"),
            parameter.get("description"),
            parameter.get("value"),
            parameter.get("unit"),
            parameter.get("range"),
            uncertainty,
            parameter.get("prior"),
            parameter.get("source"),
            parameter.get("value_type"),
            parameter.get("context"),
            evidence_text(parameter.get("evidence")),
        ]

        for index, value in enumerate(values):
            set_cell_text(row[index], value, font_size=7)

    return table


def main():
    json_files = sorted(INPUT_FOLDER.glob("*.json"))

    if not json_files:
        raise FileNotFoundError(
            f"No JSON files found in {INPUT_FOLDER.resolve()}"
        )

    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    title = document.add_heading(
        "Structured Parameter Extraction Report",
        level=0,
    )
    title.alignment = 1

    document.add_paragraph(
        "OpenAI Structured Outputs results for selected biomedical "
        "modeling papers. Missing values are shown as '-'."
    )

    document.add_heading("Papers included", level=1)

    for path in json_files:
        document.add_paragraph(path.stem, style="List Bullet")

    for file_index, json_path in enumerate(json_files):
        if file_index > 0:
            document.add_page_break()

        data = json.loads(json_path.read_text(encoding="utf-8"))
        pmid = data.get("pmid", json_path.stem)

        document.add_heading(f"PMID {pmid}", level=1)

        model = data.get("model")
        if model:
            document.add_paragraph(f"Model used: {model}")

        add_summary(document, data)

        parameters = data.get("parameters", [])

        if parameters:
            document.add_heading("Extracted records", level=2)
            add_parameter_table(document, parameters)
        else:
            document.add_paragraph(
                "No model parameter information was extracted."
            )

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_FILE)

    print(f"Created: {OUTPUT_FILE.resolve()}")


if __name__ == "__main__":
    main()
